import uvicorn
import os
from fastapi import FastAPI, UploadFile, Form, Body, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse, JSONResponse
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
import requests
from io import BytesIO
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# Import job link extractor (new module for safe job data extraction)
try:
    from job_link_extractor import extract_job_data, fetch_job_description_from_url
except ImportError as e:
    print(f"Warning: Could not import job_link_extractor: {e}")
    # Define stub functions
    def extract_job_data(url, api_key=None):
        return {"job_title": "", "job_description": "", "source": "error", "error": "Module not available"}
    def fetch_job_description_from_url(url, api_key=None):
        return ""

# Import your core functions (adjust the import path as needed)
try:
    from core import (
        parse_resume, 
        extract_job_skills, 
        analyze_skill_match, 
        generate_llm_recommendations, 
        format_for_ui_and_pdf, 
        export_to_pdf, 
        get_description_from_db, 
        calculate_ats_score
    )
except ImportError as e:
    print(f"Warning: Could not import core functions: {e}")
    # Define stub functions for testing
    def parse_resume(file): return {}
    def extract_job_skills(text): return []
    def analyze_skill_match(resume, skills): return {}
    def generate_llm_recommendations(resume, jd, match): return {}
    def format_for_ui_and_pdf(match, rec, ats): return {}
    def export_to_pdf(data): return "report.pdf"
    def get_description_from_db(role): return ""
    def calculate_ats_score(resume, jd): return {}

# Import parser functions
try:
    from parser import parse_resume_upload, parse_resume_from_path
except ImportError as e:
    print(f"Warning: Could not import parser functions: {e}")
    # Define stub functions
    async def parse_resume_upload(file): return {"error": "Parser not available"}
    def parse_resume_from_path(path): return {"error": "Parser not available"}

# Import render functions
try:
    from render import render_pdf_from_template, render_html_from_data
except ImportError as e:
    print(f"Warning: Could not import render functions: {e}")
    # Define stub functions
    def render_pdf_from_template(template, data): return JSONResponse({"error": "Render not available"})
    def render_html_from_data(data, template): return JSONResponse({"error": "Render not available"})

# Optional document extractors
try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

try:
    import docx
except Exception:
    docx = None

# Initialize FastAPI app FIRST
app = FastAPI()

# Temporary storage for the last analysis result
temp_storage = {}

# CORS middleware configuration
origins = [
    "http://localhost:5173",
    "http://localhost:8080",
    "http://localhost:8081",
    "http://localhost:3000",
    "http://localhost:8000",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:8080",
    "http://127.0.0.1:3000",
    "http://127.0.0.1:8000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Temporarily open for debugging, will refine later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    import traceback
    error_msg = str(exc)
    stack_trace = traceback.format_exc()
    print(f"ERROR: {error_msg}\n{stack_trace}")
    return JSONResponse(
        status_code=500,
        content={"success": False, "error": error_msg, "traceback": stack_trace},
    )


# Helper functions
def _extract_text_from_html(html: str) -> str:
    """Extract job description text from HTML content using regex."""
    import re

    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', html)

    # Decode HTML entities
    import html
    text = html.unescape(text)

    # Clean up whitespace
    text = re.sub(r'\s+', ' ', text).strip()

    return text


async def _extract_text_from_file(file) -> str:
    """Extract text from an uploaded file (txt, pdf, docx)."""
    if not file:
        return ""

    filename = getattr(file, 'filename', '') or ""
    content_type = getattr(file, 'content_type', '') or ""

    # Read file content
    try:
        content = await file.read()
        # reset stream so callers can read again if needed
        try:
            if hasattr(file, "file"):
                try:
                    file.file.seek(0)
                except Exception:
                    pass
        except Exception:
            pass
    except Exception as e:
        print(f"[file-extract] error reading file: {e}")
        return ""

    # Plain text
    if content_type.startswith("text/") or filename.lower().endswith(".txt"):
        try:
            return content.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            print(f"[file-extract] text decode error: {e}")
            return ""

    # PDF
    if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
        if PdfReader:
            try:
                reader = PdfReader(BytesIO(content))
                pages = []
                for p in reader.pages:
                    try:
                        pages.append(p.extract_text() or "")
                    except Exception:
                        continue
                return "\n".join(pages).strip()
            except Exception as e:
                print(f"[file-extract] pdf parse error: {e}")
                return ""
        else:
            print("[file-extract] PyPDF2 not installed")
            return ""

    # DOCX
    if (content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or
        filename.lower().endswith(".docx")):
        if docx:
            try:
                doc = docx.Document(BytesIO(content))
                full = [p.text for p in doc.paragraphs]
                return "\n".join(full).strip()
            except Exception as e:
                print(f"[file-extract] docx parse error: {e}")
                return ""
        else:
            print("[file-extract] python-docx not installed")
            return ""

    # DOC (older Word format) - fallback to text if possible
    if content_type in ["application/msword", "application/vnd.ms-word"] or filename.lower().endswith(".doc"):
        # python-docx might handle .doc files, but it's not guaranteed
        if docx:
            try:
                doc = docx.Document(BytesIO(content))
                full = [p.text for p in doc.paragraphs]
                return "\n".join(full).strip()
            except Exception as e:
                print(f"[file-extract] doc parse error: {e}")
                return ""
        else:
            print("[file-extract] python-docx not installed")
            return ""

    # Fallback: try to decode as text
    try:
        return content.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""




# NOTE: Old _fetch_job_description_from_url removed - now in job_link_extractor.py module
# Using extract_job_data() for RapidAPI + fallback HTML parsing instead


# API Endpoints

@app.get("/")
async def root():
    """Health check endpoint."""
    return {"status": "running", "message": "ELEVATR API is up!"}


@app.post("/fetch-job-description")
async def fetch_job_description_endpoint(job_description_url: str = Form(...)):
    """
    Fetch and extract job description from a URL.
    
    Uses the new job_link_extractor module:
    1. Primary: RapidAPI JSearch (safe, structured, no IP bans)
    2. Fallback: BeautifulSoup HTML parsing
    
    Response:
    {
        "success": bool,
        "text": "job description text",
        "job_title": "extracted job title (optional)",
        "message": "status message",
        "source": "api|html|error"
    }
    """
    if not job_description_url or not job_description_url.strip():
        return JSONResponse(
            {
                "success": False, 
                "text": "", 
                "job_title": "",
                "message": "No URL provided.",
                "source": "error",
            }, 
            status_code=400
        )
    
    try:
        # Use new extraction module (with RapidAPI + fallback)
        api_key = os.getenv("RAPIDAPI_KEY", "").strip()
        result = extract_job_data(job_description_url.strip(), api_key)
        
        # Check if extraction was successful
        if result.get("source") == "error":
            return JSONResponse({
                "success": False, 
                "text": "", 
                "job_title": "",
                "message": result.get("error", "Could not extract job description."),
                "source": "error",
            })
        
        # Success: return both title and description
        return {
            "success": True, 
            "text": result.get("job_description", ""),
            "job_title": result.get("job_title", ""),
            "message": f"Extracted successfully from {result.get('source', 'unknown')} source.",
            "source": result.get("source", "unknown"),
        }
        
    except Exception as e:
        print(f"Error in fetch_job_description_endpoint: {e}")
        return JSONResponse(
            {
                "success": False, 
                "text": "", 
                "job_title": "",
                "message": f"Error fetching URL: {str(e)}",
                "source": "error",
            }, 
            status_code=500
        )


class ExtractJobRequest(BaseModel):
    url: str


@app.post("/extract-job")
async def extract_job_endpoint(payload: ExtractJobRequest):
    """
    Extract job title and description from a job listing URL.
    
    Uses multi-layer extraction:
    1. RapidAPI JSearch (safe, structured)
    2. Selenium headless Chrome (JS-heavy pages)
    3. BeautifulSoup static HTML parsing
    
    Request:  {"url": "https://..."}
    Response: {"job_title": "...", "job_description": "...", "source": "api|selenium|html|error"}
    """
    if not payload.url or not payload.url.strip():
        return JSONResponse(
            {"job_title": "", "job_description": "", "source": "error",
             "error": "No URL provided."},
            status_code=400,
        )

    try:
        api_key = os.getenv("RAPIDAPI_KEY", "").strip()
        result = extract_job_data(payload.url.strip(), api_key)
        return result
    except Exception as e:
        print(f"Error in extract_job_endpoint: {e}")
        return JSONResponse(
            {"job_title": "", "job_description": "", "source": "error",
             "error": f"Extraction failed: {str(e)}"},
            status_code=500,
        )


@app.post("/analyze-resume")
async def analyze_resume(
    file: UploadFile,
    job_role: str = Form(...),
    job_description: str = Form(None),
    job_description_text: str = Form(None),
    job_description_url: str = Form(None),
    job_description_file: Optional[UploadFile] = None,
):
    """
    Analyze resume against job description.
    Priority: job_description_text > job_description_file > job_description_url > job_description > DB
    """
    # Determine final job_description_text_to_use
    jd_text = (job_description_text or "").strip()
    
    if not jd_text and job_description_file is not None:
        try:
            print(f"[jd-file] received upload filename={getattr(job_description_file, 'filename', None)} content_type={getattr(job_description_file, 'content_type', None)}")
            jd_text = await _extract_text_from_file(job_description_file)
            # Fallback: if extractor returned empty, try raw read/decode
            if not jd_text:
                try:
                    raw = await job_description_file.read()
                    jd_text = raw.decode("utf-8", errors="ignore").strip()
                    print(f"[jd-file] fallback decode length={len(jd_text)}")
                except Exception as e:
                    print(f"[jd-file] fallback read error: {e}")
        except Exception as e:
            print(f"[jd-file] extraction error: {e}")
            jd_text = ""

        if jd_text:
            print(f"[jd-file] extracted length={len(jd_text)} preview={jd_text[:200].replace('\n',' ')}")
    
    if not jd_text and job_description_url:
        # Use new job_link_extractor module (RapidAPI + fallback parsing)
        api_key = os.getenv("RAPIDAPI_KEY", "").strip()
        jd_text = fetch_job_description_from_url(job_description_url, api_key)
    
    if not jd_text and job_description:
        jd_text = job_description

    # Fallback to DB if still empty
    if not jd_text:
        jd_text = get_description_from_db(job_role)

    resume_data = parse_resume(file)
    required_skills = extract_job_skills(jd_text)
    match_info = analyze_skill_match(resume_data, required_skills)
    ats_data = calculate_ats_score(resume_data, jd_text)
    recommendations = generate_llm_recommendations(resume_data, jd_text, match_info)
    formatted = format_for_ui_and_pdf(match_info, recommendations, ats_data)

    temp_storage["last_result"] = formatted

    return {"result": formatted}


@app.options("/export-pdf")
async def export_pdf_options():
    """CORS preflight for export-pdf endpoint."""
    return {"message": "OK"}


@app.get("/export-pdf")
async def export_pdf():
    """Export the last analysis result as a PDF."""
    formatted_data = temp_storage.get("last_result")
    if not formatted_data:
        return JSONResponse(
            {"error": "No analysis result available to export."}, 
            status_code=404
        )

    pdf_path = export_to_pdf(formatted_data)
    return FileResponse(
        path=pdf_path,
        filename="ResumeReport(ELEVATR).pdf",
        media_type="application/pdf"
    )


@app.get("/templates")
async def list_templates():
    """List available LaTeX templates."""
    candidates: List[str] = []
    root_templates = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
    backend_templates = os.path.join(os.path.dirname(__file__), "templates")
    
    for d in [backend_templates, root_templates]:
        if os.path.isdir(d):
            for f in os.listdir(d):
                if f.endswith(".tex"):
                    candidates.append(f)
    
    # De-duplicate while preserving order
    seen = set()
    unique = [x for x in candidates if not (x in seen or seen.add(x))]
    return {"templates": unique or ["modern.tex", "classic.tex", "professional.tex"]}


@app.post("/upload_resume")
async def upload_resume(file: UploadFile):
    """Parse an uploaded resume file and extract information."""
    try:
        result = await parse_resume_upload(file)
        return result
    except Exception as e:
        print(f"Error in upload_resume: {e}")
        return JSONResponse(
            {"error": f"Failed to parse resume: {str(e)}", "parsed": None},
            status_code=500
        )


@app.post("/upload_resume_path")
async def upload_resume_path(file_path: str = Body(..., embed=True)):
    """Parse a resume from a file path and extract information."""
    try:
        result = parse_resume_from_path(file_path)
        return result
    except Exception as e:
        print(f"Error in upload_resume_path: {e}")
        return JSONResponse(
            {"error": f"Failed to parse resume: {str(e)}", "parsed": None},
            status_code=500
        )


@app.post("/generate_resume")
async def generate_resume(
    template: str = Body(...),
    data: dict = Body(...),
    format: str = Body(...),
):
    """Generate a resume in PDF or HTML format from data and template."""
    try:
        if format.lower() == "pdf":
            # Return PDF as streaming response
            return render_pdf_from_template(template, data)
        elif format.lower() == "html":
            # Return HTML response
            return render_html_from_data(data, template)
        else:
            return JSONResponse(
                {"error": f"Unsupported format: {format}. Use 'pdf' or 'html'."},
                status_code=400
            )
    except Exception as e:
        print(f"Error in generate_resume: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"error": f"Failed to generate resume: {str(e)}"},
            status_code=500
        )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print(f"Starting ELEVATR API server on port {port}...")
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)