from typing import Dict, Any, List, Optional, Tuple
from fastapi import UploadFile
import re
import io
import os
from collections import defaultdict

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import spacy
    _NLP = spacy.load("en_core_web_md")
    SPACY_AVAILABLE = True
except Exception:
    _NLP = None
    SPACY_AVAILABLE = False


# Enhanced keyword lists
DEGREE_KEYWORDS = [
    "bachelor", "master", "phd", "doctorate", "b.sc", "m.sc", "b.tech", "m.tech",
    "b.e.", "m.e.", "mba", "bachelors", "masters", "bba", "mca", "bca",
    "associate", "diploma", "certification", "ph.d", "m.d", "j.d", "ll.b", "ll.m"
]

# More comprehensive date patterns
DATE_PATTERNS = [
    r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}',
    r'\d{1,2}/\d{4}',
    r'\d{4}\s*-\s*\d{4}',
    r'\d{4}\s*–\s*\d{4}',
    r'\d{4}\s*to\s*\d{4}',
    r'\d{4}\s*-\s*(?:Present|Current|Now)',
    r'\d{4}',
]

BLACKLIST_TOKENS = {
    'resume', 'cv', 'curriculum', 'vitae', 'portfolio', 'streamlit', 'github', 'linkedin', 
    'email', 'phone', 'contact', 'profile', 'summary', 'objective', 'skills', 'experience',
    'education', 'projects', 'certifications', 'references', 'available', 'upon', 'request'
}

# Expanded tech skills database
TECH_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'rust',
    'swift', 'kotlin', 'php', 'scala', 'r', 'matlab', 'perl', 'shell', 'bash',
    
    # Web Technologies
    'react', 'angular', 'vue', 'vue.js', 'node.js', 'express', 'django', 'flask',
    'spring', 'fastapi', 'nest.js', 'next.js', 'svelte', 'html', 'css', 'sass',
    'less', 'bootstrap', 'tailwind', 'webpack', 'vite',
    
    # Databases
    'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'cassandra', 'dynamodb',
    'oracle', 'sqlite', 'mariadb', 'elasticsearch', 'neo4j',
    
    # Cloud & DevOps
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'circleci', 'gitlab',
    'terraform', 'ansible', 'chef', 'puppet', 'ci/cd', 'microservices',
    
    # Data Science & ML
    'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'matplotlib',
    'seaborn', 'jupyter', 'spark', 'hadoop', 'airflow', 'kafka', 'nlp', 'computer vision',
    
    # Mobile
    'react native', 'flutter', 'android', 'ios', 'xamarin',
    
    # Other
    'git', 'linux', 'unix', 'windows', 'agile', 'scrum', 'jira', 'rest', 'graphql',
    'api', 'microservices', 'blockchain', 'web3'
}

SOFT_SKILLS = {
    'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
    'project management', 'time management', 'collaboration', 'adaptability', 'creativity',
    'analytical', 'presentation', 'negotiation', 'mentoring', 'strategic thinking'
}


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF with better handling"""
    if not PYMUPDF_AVAILABLE:
        return "PDF parsing not available - PyMuPDF not installed"
    
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Try different extraction methods for better accuracy
            text = page.get_text("text")
            if not text.strip():
                # Fallback to blocks if text extraction fails
                text = page.get_text("blocks")
            text_parts.append(text if isinstance(text, str) else "")
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"


def extract_name_by_font_from_pdf(content: bytes) -> str:
    """Enhanced font-based name extraction with better heuristics"""
    if not PYMUPDF_AVAILABLE:
        return ""
    
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        candidates = []
        
        # Scan first page only for name
        if doc.page_count > 0:
            page = doc.load_page(0)
            text_dict = page.get_text("dict")
            
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:  # Only text blocks
                    continue
                    
                for line in block.get("lines", []):
                    # Check if line is in the top 25% of the page
                    bbox = line.get("bbox", [0, 0, 0, 0])
                    if bbox[1] > page.rect.height * 0.25:
                        continue
                    
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        size = float(span.get("size", 0))
                        
                        if text and size >= 12 and 5 <= len(text) <= 60:
                            # Check if text looks like a name
                            if _looks_like_name_advanced(text):
                                candidates.append((size, text, bbox[1]))
        
        doc.close()
        
        if candidates:
            # Sort by: size (desc), then position (asc)
            candidates.sort(key=lambda x: (-x[0], x[2]))
            return candidates[0][1]
        
        return ""
    except Exception:
        return ""


def extract_text_from_docx(content: bytes) -> str:
    """Enhanced DOCX extraction with better fallback handling"""
    if not DOCX_AVAILABLE:
        try:
            import tempfile
            import docx2txt
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
                tmp.write(content)
                tmp.flush()
                extracted = docx2txt.process(tmp.name) or ""
                return extracted
        except Exception:
            return "DOCX parsing not available - python-docx not installed"

    try:
        doc = Document(io.BytesIO(content))
        lines: List[str] = []
        
        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        
        # Extract tables
        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(par.text.strip() for par in cell.paragraphs if par.text.strip())
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    lines.append(" | ".join(row_text))
        
        # Extract headers and footers
        for section in doc.sections:
            try:
                if section.header:
                    for p in section.header.paragraphs:
                        if p.text.strip():
                            lines.append(p.text.strip())
                if section.footer:
                    for p in section.footer.paragraphs:
                        if p.text.strip():
                            lines.append(p.text.strip())
            except Exception:
                pass
        
        text = "\n".join(lines)
        if text.strip():
            return text
        
        raise ValueError("Empty text from python-docx")
        
    except Exception:
        # Fallback to docx2txt
        try:
            import tempfile
            import docx2txt
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
                tmp.write(content)
                tmp.flush()
                extracted = docx2txt.process(tmp.name) or ""
                return extracted if extracted else "Error parsing DOCX: empty after fallback"
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}"


def extract_email(text: str) -> str:
    """Enhanced email extraction"""
    email_pattern = r'\b[A-Za-z0-9][A-Za-z0-9._%+-]*@[A-Za-z0-9][A-Za-z0-9.-]*\.[A-Z|a-z]{2,}\b'
    matches = re.findall(email_pattern, text)
    
    if matches:
        # Prioritize non-example emails
        for email in matches:
            if not any(x in email.lower() for x in ['example', 'domain', 'email', 'test']):
                return email
        return matches[0]
    
    return ""


def extract_phone(text: str) -> str:
    """Enhanced phone extraction with international support"""
    phone_patterns = [
        r'\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',  # International
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
        r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # Simple US
        r'\d{10}',  # 10 digits
        r'\d{5}[-.\s]?\d{5}',  # India format
    ]
    
    for pattern in phone_patterns:
        match = re.search(pattern, text)
        if match:
            phone = match.group(0)
            # Validate it's actually a phone number (not just random digits)
            if len(re.findall(r'\d', phone)) >= 10:
                return phone
    
    return ""


def _looks_like_name_advanced(text: str) -> bool:
    """Advanced name validation with better heuristics"""
    text = text.strip()
    
    # Remove common prefixes/suffixes
    text = re.sub(r'^(Mr\.?|Mrs\.?|Ms\.?|Dr\.?|Prof\.?)\s+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+(Jr\.?|Sr\.?|II|III|IV)$', '', text, flags=re.IGNORECASE)
    
    parts = [p for p in re.split(r'\s+', text) if p]
    
    if not (1 <= len(parts) <= 5):
        return False
    
    # Check for blacklisted tokens
    for part in parts:
        if part.lower() in BLACKLIST_TOKENS or part.lower() in TECH_SKILLS:
            return False
    
    # Must have at least one capitalized word
    capitalized_count = sum(1 for p in parts if p and p[0].isupper())
    if capitalized_count == 0:
        return False
    
    # All parts should be alphabetic (with hyphens/apostrophes allowed)
    for part in parts:
        if not re.match(r"^[A-Za-z]([A-Za-z\-'\.]*[A-Za-z])?$", part):
            return False
    
    # Check length constraints
    if len(text) < 4 or len(text) > 60:
        return False
    
    # At least one part should be longer than 1 character
    if all(len(p) <= 1 for p in parts):
        return False
    
    return True


def extract_name_multimethod(text: str, pdf_content: Optional[bytes] = None) -> str:
    """Multi-method name extraction with scoring"""
    candidates = []
    
    # Method 1: Font-based (for PDFs)
    if pdf_content:
        font_name = extract_name_by_font_from_pdf(pdf_content)
        if font_name:
            candidates.append(('font', font_name, 10))
    
    # Method 2: spaCy NER
    if SPACY_AVAILABLE and text:
        try:
            doc = _NLP(text[:5000])  # Limit to first 5000 chars for performance
            for ent in doc.ents:
                if ent.label_ == "PERSON" and _looks_like_name_advanced(ent.text):
                    candidates.append(('spacy', ent.text, 8))
                    break
        except Exception:
            pass
    
    # Method 3: Header heuristic (first qualifying line)
    lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
    for i, line in enumerate(lines[:15]):
        if _looks_like_name_advanced(line):
            score = 7 - (i * 0.3)  # Earlier lines get higher scores
            candidates.append(('header', line, score))
            break
    
    # Method 4: Pattern matching for "Name: John Doe" format
    name_patterns = [
        r'(?:Name|Full Name|Candidate)\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
    ]
    for pattern in name_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            name = match.group(1)
            if _looks_like_name_advanced(name):
                candidates.append(('pattern', name, 6))
                break
    
    # Sort by score and return best candidate
    if candidates:
        candidates.sort(key=lambda x: -x[2])
        return candidates[0][1]
    
    return ""


def extract_location(text: str) -> str:
    """Enhanced location extraction"""
    # Pattern 1: City, State/Country format
    location_patterns = [
        r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z]{2}|[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b',
        r'\b([A-Z][a-z]+),\s*([A-Z][a-z]+)\b',
    ]
    
    for pattern in location_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Filter out obvious non-locations
            for match in matches:
                location = f"{match[0]}, {match[1]}"
                if not any(word.lower() in BLACKLIST_TOKENS for word in match):
                    return location
    
    # Pattern 2: Look for location keywords
    location_keywords = ['location', 'address', 'based in', 'residing in']
    for keyword in location_keywords:
        pattern = rf'{keyword}\s*[:\-]?\s*([A-Z][a-z]+(?:[\s,]+[A-Z][a-z]+)*)'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return ""


def extract_linkedin(text: str) -> str:
    """Enhanced LinkedIn extraction"""
    patterns = [
        r'https?://(?:www\.)?linkedin\.com/in/[\w\-]+/?',
        r'linkedin\.com/in/[\w\-]+/?',
        r'(?:LinkedIn|Linkedin|linkedin)\s*[:\-]?\s*(?:https?://)?(?:www\.)?linkedin\.com/in/([\w\-]+)/?',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            url = match.group(0)
            if not url.startswith('http'):
                url = 'https://' + url.lstrip()
            # Clean up the URL
            url = re.sub(r'/$', '', url)
            return url
    
    return ""


def extract_github(text: str) -> str:
    """Enhanced GitHub extraction"""
    patterns = [
        r'https?://(?:www\.)?github\.com/[\w\-]+/?',
        r'github\.com/[\w\-]+/?',
        r'(?:GitHub|Github|github)\s*[:\-]?\s*(?:https?://)?(?:www\.)?github\.com/([\w\-]+)/?',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            url = match.group(0)
            if not url.startswith('http'):
                url = 'https://' + url.lstrip()
            url = re.sub(r'/$', '', url)
            return url
    
    return ""


def extract_website(text: str) -> str:
    """Extract personal website/portfolio"""
    # Exclude common social media and email domains
    exclude_domains = ['linkedin', 'github', 'facebook', 'twitter', 'instagram', 'gmail', 'yahoo', 'outlook']
    
    pattern = r'(?:https?://)?(?:www\.)?[\w\-]+\.[\w\-]+(?:\.[\w\-]+)*/?[\w\-]*'
    matches = re.findall(pattern, text)
    
    for match in matches:
        if not any(domain in match.lower() for domain in exclude_domains):
            if not match.startswith('http'):
                match = 'https://' + match
            return match
    
    return ""


def extract_skills_enhanced(text: str) -> List[Dict[str, str]]:
    """Enhanced skill extraction with categorization"""
    text_lower = text.lower()
    found_skills = defaultdict(list)
    
    # Extract technical skills
    for skill in TECH_SKILLS:
        # Use word boundaries for better matching
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found_skills['technical'].append(skill.title())
    
    # Extract soft skills
    for skill in SOFT_SKILLS:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found_skills['soft'].append(skill.title())
    
    # Look for skills section and extract additional skills
    skills_section_pattern = r'(?:skills|technical skills|core competencies|technologies)[:\s]+(.*?)(?:\n\n|\n[A-Z][A-Z\s]+:)'
    match = re.search(skills_section_pattern, text, re.IGNORECASE | re.DOTALL)
    if match:
        skills_text = match.group(1)
        # Extract comma or pipe separated skills
        additional_skills = re.split(r'[,|\n•·]', skills_text)
        for skill in additional_skills:
            skill = skill.strip()
            if skill and 2 <= len(skill) <= 30 and not any(x in skill.lower() for x in BLACKLIST_TOKENS):
                if skill.lower() not in [s.lower() for s in found_skills['technical'] + found_skills['soft']]:
                    found_skills['technical'].append(skill)
    
    # Convert to expected format
    result = []
    for skill in found_skills['technical'][:30]:  # Limit to 30 skills
        result.append({
            "name": skill,
            "category": "technical",
            "proficiency": "intermediate"
        })
    for skill in found_skills['soft'][:10]:
        result.append({
            "name": skill,
            "category": "soft",
            "proficiency": "intermediate"
        })
    
    return result


def extract_section_content(text: str, keywords: List[str], next_section_keywords: List[str] = None) -> str:
    """Extract content from a specific section"""
    for keyword in keywords:
        pattern = rf'(?:{keyword})[:\s]*\n(.*?)(?=\n\n[A-Z]|$)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            content = match.group(1).strip()
            # If next section keywords provided, truncate at their occurrence
            if next_section_keywords:
                for next_keyword in next_section_keywords:
                    next_pattern = rf'\n\s*{next_keyword}\s*[:\n]'
                    next_match = re.search(next_pattern, content, re.IGNORECASE)
                    if next_match:
                        content = content[:next_match.start()]
                        break
            return content
    return ""


def parse_date_range(text: str) -> Tuple[str, str]:
    """Parse date ranges like 'Jan 2020 - Present' or '2019-2021'"""
    # Pattern for month year ranges
    pattern1 = r'([A-Z][a-z]+\s+\d{4})\s*[-–to]\s*([A-Z][a-z]+\s+\d{4}|Present|Current|Now)'
    match = re.search(pattern1, text, re.IGNORECASE)
    if match:
        return match.group(1), match.group(2)
    
    # Pattern for year ranges
    pattern2 = r'(\d{4})\s*[-–to]\s*(\d{4}|Present|Current|Now)'
    match = re.search(pattern2, text, re.IGNORECASE)
    if match:
        return match.group(1), match.group(2)
    
    # Single date
    for pattern in DATE_PATTERNS:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0), ""
    
    return "", ""


def extract_work_experience_enhanced(text: str) -> List[Dict[str, Any]]:
    """Enhanced work experience extraction"""
    experience_keywords = ['experience', 'work experience', 'employment', 'work history', 
                          'professional experience', 'career history']
    education_keywords = ['education', 'academic', 'qualifications']
    
    # Find the experience section
    exp_section = extract_section_content(text, experience_keywords, education_keywords + ['projects', 'skills'])
    if not exp_section:
        return []
    
    experiences = []
    
    # Split by job entries (usually separated by blank lines or clear delimiters)
    # Look for patterns like "Job Title at Company"
    job_pattern = r'([A-Z][A-Za-z\s&,]+?)(?:\s+(?:at|@|-|,)\s+)([A-Z][A-Za-z\s&,.]+?)(?:\n|\s{2,})'
    
    entries = re.split(r'\n\s*\n', exp_section)
    
    for entry in entries:
        if not entry.strip() or len(entry.strip()) < 20:
            continue
        
        # Try to extract job title and company
        match = re.search(job_pattern, entry)
        if match:
            title = match.group(1).strip()
            company = match.group(2).strip()
        else:
            # Fallback: use first line as title
            lines = entry.strip().split('\n')
            title = lines[0].strip() if lines else ""
            company = ""
        
        # Extract dates
        start_date, end_date = parse_date_range(entry)
        
        # Extract description (remaining text after title/company/dates)
        description = entry.strip()
        # Remove title, company, dates from description
        for item in [title, company, start_date, end_date]:
            if item:
                description = description.replace(item, '')
        description = re.sub(r'\s+', ' ', description).strip()
        
        experiences.append({
            'title': title,
            'company': company,
            'startDate': start_date,
            'endDate': end_date or "Present",
            'description': description[:500]  # Limit description length
        })
    
    return experiences


def extract_education_enhanced(text: str) -> List[Dict[str, Any]]:
    """Enhanced education extraction"""
    education_keywords = ['education', 'academic background', 'qualifications', 'degrees']
    experience_keywords = ['experience', 'work experience', 'employment']
    
    # Find the education section
    edu_section = extract_section_content(text, education_keywords, experience_keywords + ['projects', 'skills'])
    if not edu_section:
        return []
    
    education = []
    
    # Split into entries
    entries = re.split(r'\n\s*\n', edu_section)
    
    for entry in entries:
        if not entry.strip() or len(entry.strip()) < 10:
            continue
        
        # Extract degree
        degree = ""
        for keyword in DEGREE_KEYWORDS:
            pattern = rf'\b{keyword}[a-z]*(?:\s+of\s+[A-Za-z\s]+)?'
            match = re.search(pattern, entry, re.IGNORECASE)
            if match:
                degree = match.group(0).strip()
                break
        
        # Extract institution
        institution = ""
        # Look for capitalized multi-word phrases (likely institution names)
        inst_pattern = r'([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)+(?:\s+(?:University|College|Institute|School|Academy))?)'
        inst_match = re.search(inst_pattern, entry)
        if inst_match:
            institution = inst_match.group(1).strip()
        
        # Extract dates
        start_date, end_date = parse_date_range(entry)
        
        # Extract GPA if present
        gpa = ""
        gpa_pattern = r'GPA[:\s]*(\d+\.?\d*(?:/\d+\.?\d*)?)'
        gpa_match = re.search(gpa_pattern, entry, re.IGNORECASE)
        if gpa_match:
            gpa = gpa_match.group(1)
        
        education.append({
            'degree': degree or entry.split('\n')[0].strip(),
            'institution': institution,
            'startDate': start_date,
            'endDate': end_date,
            'gpa': gpa,
            'description': entry.strip()[:300]
        })
    
    return education


def extract_projects(text: str) -> List[Dict[str, Any]]:
    """Extract project information"""
    project_keywords = ['projects', 'personal projects', 'key projects', 'notable projects']
    next_keywords = ['experience', 'education', 'skills', 'certifications']
    
    project_section = extract_section_content(text, project_keywords, next_keywords)
    if not project_section:
        return []
    
    projects = []
    entries = re.split(r'\n\s*\n', project_section)
    
    for entry in entries:
        if not entry.strip() or len(entry.strip()) < 15:
            continue
        
        lines = entry.strip().split('\n')
        title = lines[0].strip()
        
        # Extract technologies used
        tech_pattern = r'(?:Technologies|Tech Stack|Built with)[:\s]*(.*?)(?:\n|$)'
        tech_match = re.search(tech_pattern, entry, re.IGNORECASE)
        technologies = tech_match.group(1).strip() if tech_match else ""
        
        # Extract URL if present
        url = ""
        url_pattern = r'(?:URL|Link|Demo|GitHub)[:\s]*(https?://[^\s]+)'
        url_match = re.search(url_pattern, entry, re.IGNORECASE)
        if url_match:
            url = url_match.group(1)
        
        description = entry.strip()
        
        projects.append({
            'title': title,
            'description': description[:400],
            'technologies': technologies,
            'url': url
        })
    
    return projects


def extract_certifications(text: str) -> List[Dict[str, str]]:
    """Extract certifications"""
    cert_keywords = ['certifications', 'certificates', 'licenses', 'credentials']
    next_keywords = ['experience', 'education', 'skills', 'projects']
    
    cert_section = extract_section_content(text, cert_keywords, next_keywords)
    if not cert_section:
        return []
    
    certifications = []
    lines = [l.strip() for l in cert_section.split('\n') if l.strip()]
    
    for line in lines:
        if len(line) < 5:
            continue
        
        # Extract date if present
        date_match = None
        for pattern in DATE_PATTERNS:
            date_match = re.search(pattern, line, re.IGNORECASE)
            if date_match:
                break
        
        cert_name = line
        issue_date = ""
        
        if date_match:
            issue_date = date_match.group(0)
            cert_name = line.replace(issue_date, '').strip()
        
        # Clean up cert name
        cert_name = re.sub(r'^[•\-\*]\s*', '', cert_name).strip()
        
        certifications.append({
            'name': cert_name,
            'issuer': '',
            'date': issue_date
        })
    
    return certifications


def extract_summary(text: str, name: str) -> str:
    """Extract professional summary or objective"""
    summary_keywords = ['summary', 'professional summary', 'objective', 'profile', 
                        'about me', 'career objective', 'professional profile']
    
    # Try to find summary section
    summary = extract_section_content(text, summary_keywords, 
                                     ['experience', 'education', 'skills', 'projects'])
    if summary:
        return summary[:500]
    
    # Fallback: look for paragraph after name/contact info
    lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
    in_contact_section = False
    for i, line in enumerate(lines[:30]):
        # Skip contact info lines
        if any(x in line.lower() for x in ['email', 'phone', '@', 'linkedin', 'github', 'http']):
            in_contact_section = True
            continue
        
        # If we found a substantial line after contact info
        if in_contact_section and len(line) > 50 and not line.isupper():
            # Make sure it's not a section header
            if not any(keyword in line.lower() for keyword in ['experience', 'education', 'skills', 'projects']):
                return line
    
    return ""


async def parse_resume_upload(file: UploadFile) -> Dict[str, Any]:
    """Enhanced resume parser with significantly improved accuracy"""
    content = await file.read()
    filename = file.filename or ""
    
    # Extract text based on file type
    if filename.lower().endswith('.pdf'):
        raw_text = extract_text_from_pdf(content)
        pdf_content = content
    elif filename.lower().endswith(('.doc', '.docx')):
        raw_text = extract_text_from_docx(content)
        pdf_content = None
    else:
        return {
            "error": f"Unsupported file type: {filename}",
            "parsed": None,
            "rawText": "",
            "confidence": 0.0
        }
    
    # Check if extraction was successful
    if raw_text.startswith("Error") or len(raw_text.strip()) < 50:
        return {
            "error": "Failed to extract text from resume",
            "parsed": None,
            "rawText": raw_text,
            "confidence": 0.0
        }
    
    # Extract all information
    name = extract_name_multimethod(raw_text, pdf_content)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    location = extract_location(raw_text)
    linkedin = extract_linkedin(raw_text)
    github = extract_github(raw_text)
    website = extract_website(raw_text)
    summary = extract_summary(raw_text, name)
    skills = extract_skills_enhanced(raw_text)
    work_experience = extract_work_experience_enhanced(raw_text)
    education = extract_education_enhanced(raw_text)
    projects = extract_projects(raw_text)
    certifications = extract_certifications(raw_text)
    
    # Calculate confidence score
    confidence = 0.0
    if name: confidence += 0.25
    if email: confidence += 0.20
    if phone: confidence += 0.10
    if work_experience: confidence += 0.20
    if education: confidence += 0.15
    if skills: confidence += 0.10
    
    confidence = min(confidence, 1.0)
    
    return {
        "parsed": {
            "personalInfo": {
                "name": name,
                "email": email,
                "phone": phone,
                "location": location,
                "linkedin": linkedin,
                "github": github,
                "website": website,
            },
            "summary": summary,
            "education": education,
            "workExperience": work_experience,
            "projects": projects,
            "skills": skills,
            "certifications": certifications,
        },
        "rawText": raw_text,
        "confidence": confidence,
    }


def parse_resume_from_path(file_path: str) -> Dict[str, Any]:
    """Synchronous version for parsing from file path"""
    if not os.path.isfile(file_path):
        return {"error": f"File not found: {file_path}"}
    
    try:
        with open(file_path, "rb") as f:
            content = f.read()
    except Exception as e:
        return {"error": f"Failed to read file: {e}"}
    
    lower = file_path.lower()
    if lower.endswith('.pdf'):
        raw_text = extract_text_from_pdf(content)
        pdf_content = content
    elif lower.endswith(('.doc', '.docx')):
        raw_text = extract_text_from_docx(content)
        pdf_content = None
    else:
        return {"error": f"Unsupported file type: {file_path}"}
    
    if raw_text.startswith("Error") or len(raw_text.strip()) < 50:
        return {
            "error": "Failed to extract text from resume",
            "parsed": None,
            "rawText": raw_text,
            "confidence": 0.0
        }
    
    # Extract all information (same as async version)
    name = extract_name_multimethod(raw_text, pdf_content)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    location = extract_location(raw_text)
    linkedin = extract_linkedin(raw_text)
    github = extract_github(raw_text)
    website = extract_website(raw_text)
    summary = extract_summary(raw_text, name)
    skills = extract_skills_enhanced(raw_text)
    work_experience = extract_work_experience_enhanced(raw_text)
    education = extract_education_enhanced(raw_text)
    projects = extract_projects(raw_text)
    certifications = extract_certifications(raw_text)
    
    # Calculate confidence
    confidence = 0.0
    if name: confidence += 0.25
    if email: confidence += 0.20
    if phone: confidence += 0.10
    if work_experience: confidence += 0.20
    if education: confidence += 0.15
    if skills: confidence += 0.10
    confidence = min(confidence, 1.0)
    
    return {
        "parsed": {
            "personalInfo": {
                "name": name,
                "email": email,
                "phone": phone,
                "location": location,
                "linkedin": linkedin,
                "github": github,
                "website": website,
            },
            "summary": summary,
            "education": education,
            "workExperience": work_experience,
            "projects": projects,
            "skills": skills,
            "certifications": certifications,
        },
        "rawText": raw_text,
        "confidence": confidence,
    }